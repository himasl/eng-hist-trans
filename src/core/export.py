import io

from docx import Document


def text_to_docx_bytes(text: str, title: str = "Транскрипция") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.splitlines():
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
